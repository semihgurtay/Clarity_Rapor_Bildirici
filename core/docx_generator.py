import os
import re
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

log = logging.getLogger("DocxGenerator")

def set_cell_background(cell, hex_color):
    """Word tablosundaki hücrenin arka plan rengini ayarla."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Hücre padding/margin değerlerini ayarla."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_to_docx(markdown_text: str, docx_path: str):
    """Markdown formatındaki metni profesyonelce biçimlendirilmiş Word belgesine dönüştür."""
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla (2.5 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Varsayılan stilleri düzenle
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.2
    style_normal.paragraph_format.space_after = Pt(6)

    lines = markdown_text.splitlines()

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
            
        # Sadece separator satırlarını temizle
        rows = [r for r in table_rows if not re.match(r'^[\s|:-]+$', "".join(r))]
        if not rows:
            table_rows = []
            in_table = False
            return

        # Word tablosu oluştur
        col_count = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=col_count)
        t.style = 'Table Grid'
        
        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                if c_idx >= col_count:
                    continue
                cell = t.cell(r_idx, c_idx)
                cell.text = val.strip()
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                # İlk satır (Başlık) stili
                if r_idx == 0:
                    set_cell_background(cell, '0F3460') # Lacivert
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    # Zebra şeritleri
                    if r_idx % 2 == 0:
                        set_cell_background(cell, 'F4F6F9')
                    p = cell.paragraphs[0]
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)

        # Tablo sonrası boşluk bırak
        doc.add_paragraph()
        table_rows = []
        in_table = False

    for line_raw in lines:
        line = line_raw.strip()
        
        # Boş satır kontrolü
        if not line:
            if in_table:
                flush_table()
            continue

        # Tablo satırı tespiti
        if line.startswith('|'):
            if not in_table:
                in_table = True
            parts = [p.strip() for p in line.split('|')[1:-1]]
            table_rows.append(parts)
            continue
        elif in_table:
            flush_table()

        # Alert / Blockquote
        if line.startswith('>'):
            content = line.replace('>', '', 1).strip()
            content = re.sub(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', r'[\1]', content)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(content)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
            pPr.append(shd)
            continue

        # Başlıklar
        if line.startswith('# '):
            title_text = line[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            run = p.add_run(title_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="8" w:color="E94560"/></w:pBdr>')
            pPr.append(pBdr)
            continue
            
        elif line.startswith('## '):
            h_text = line[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            
            run = p.add_run(h_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
            continue
            
        elif line.startswith('### '):
            h_text = line[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(h_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            continue

        # Yatay çizgi (---)
        if re.match(r'^[-*_]{3,}$', line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            pPr.append(pBdr)
            continue

        # Liste Elemanları
        if line.startswith(('- ', '* ')):
            list_text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            
            parts = re.split(r'(\*\*.*?\*\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
            continue

        # Numaralı Liste Elemanları
        num_match = re.match(r'^(\d+)\.\s(.*)', line)
        if num_match:
            list_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            
            parts = re.split(r'(\*\*.*?\*\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
            continue

        # Normal Paragraf
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            else:
                p.add_run(part)

    if in_table:
        flush_table()

    doc.save(docx_path)
    log.info(f"Word belgesi başarıyla oluşturuldu: {docx_path}")
